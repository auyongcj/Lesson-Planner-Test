from __future__ import annotations

from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .schema import ScheduleDoc
from .lesson_schema import LessonPlanDoc


def _set_cell_border(cell, **kwargs):
    """
    kwargs keys: top, left, bottom, right
    each value is dict: {"sz": 12, "val": "single", "color": "000000"}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            element = tcBorders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            element.set(qn("w:val"), edge_data.get("val", "single"))
            element.set(qn("w:sz"), str(edge_data.get("sz", 8)))
            element.set(qn("w:color"), edge_data.get("color", "000000"))


def export_training_programme_docx_bytes(doc: ScheduleDoc) -> bytes:
    d = Document()

    # base font
    style = d.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Titles
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc.programme_title)
    run.bold = True
    run.underline = True

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc.unit_title)
    run.bold = True

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{doc.programme_title} {doc.programme_date_line}")

    d.add_paragraph()  # spacer

    # Table
    headers = ["Time", "Activity", "Venue", "IC", "Attire", "Remarks"]
    table = d.add_table(rows=1, cols=len(headers))
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        para = hdr_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        r.bold = True

        # header: no top border, thick bottom, verticals
        _set_cell_border(
            hdr_cells[i],
            top={"val": "nil"},
            bottom={"sz": 16},
            left={"sz": 8},
            right={"sz": 8},
        )

    for r in doc.rows:
        row_cells = table.add_row().cells
        vals = [r.time, r.activity, r.venue, r.ic, r.attire, r.remarks]
        for i, v in enumerate(vals):
            para = row_cells[i].paragraphs[0]
            # top aligned feel: just left align remarks, center others
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 5 else WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(v or "")

            # body: vertical borders only, no horizontal
            _set_cell_border(
                row_cells[i],
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"sz": 8},
                right={"sz": 8},
            )

    # outer border: easiest is set bottom border on last row
    last_row = table.rows[-1]
    for c in last_row.cells:
        _set_cell_border(c, bottom={"sz": 16})

    d.add_paragraph()
    d.add_paragraph("Notes:").runs[0].bold = True
    for n in doc.notes:
        d.add_paragraph(n, style="List Number")

    d.add_paragraph()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Authenticated by: {doc.authenticated_by}")

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Vetted by: {doc.vetted_by}")

    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


def export_lesson_plan_docx_bytes(lesson: LessonPlanDoc) -> bytes:
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    def heading(text: str):
        p = d.add_paragraph()
        r = p.add_run(text)
        r.bold = True

    heading("1. LESSON TITLE")
    d.add_paragraph(lesson.lesson_title or "")

    heading("2. LESSON OBJECTIVES")
    d.add_paragraph("At the end of the lesson, participants are expected to:")
    for x in lesson.lesson_objectives:
        d.add_paragraph(x, style="List Bullet")

    heading("3. PRIOR KNOWLEDGE")
    for x in (lesson.prior_knowledge or ["NIL."]):
        d.add_paragraph(x, style="List Bullet")

    # You can extend this to match your sample tables exactly.
    # For now, export key sections and Method remarks structure.
    heading("6. METHOD OF INSTRUCTION")
    d.add_paragraph(f"TRAINING PERSONNEL: {lesson.training_personnel_line}")

    for mr in lesson.method_rows:
        d.add_paragraph(f"Stage {mr.stage}: {mr.activity}", style=None)
        rb = mr.remarks
        if rb.lead:
            d.add_paragraph(rb.lead)
        # list style
        for item in rb.items:
            d.add_paragraph(item.text, style="List Bullet")
            for sub in item.subitems:
                d.add_paragraph(sub, style="List Number")  # simple; can be roman with custom style

    heading("10. AUTHENTICATION")
    d.add_paragraph(f"Prepared by: {lesson.prepared_by}")
    d.add_paragraph(f"Vetted by: {lesson.vetted_by}")
    d.add_paragraph(f"Date: {lesson.date_text}")

    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()